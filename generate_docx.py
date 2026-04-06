from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


MONTH_NAMES = {
    "january",
    "february",
    "march",
    "april",
    "may",
    "june",
    "july",
    "august",
    "september",
    "october",
    "november",
    "december",
}


def is_upper_heading(line: str) -> bool:
    stripped = line.strip()
    if not stripped:
        return False
    if stripped.startswith("-"):
        return False
    letters = [char for char in stripped if char.isalpha()]
    return bool(letters) and all(char.isupper() for char in letters)


def looks_like_role_or_company(line: str) -> bool:
    return " | " in line


def looks_like_date_or_subrole(line: str) -> bool:
    lowered = line.lower()
    if " | " in line:
        return False
    return any(month in lowered for month in MONTH_NAMES) and "-" in line


def ensure_styles(document: Document) -> None:
    styles = document.styles

    if "Resume Section" not in styles:
        section_style = styles.add_style("Resume Section", WD_STYLE_TYPE.PARAGRAPH)
        section_style.font.name = "Arial"
        section_style.font.size = Pt(11)
        section_style.font.bold = True

    if "Resume Role" not in styles:
        role_style = styles.add_style("Resume Role", WD_STYLE_TYPE.PARAGRAPH)
        role_style.font.name = "Arial"
        role_style.font.size = Pt(10.5)
        role_style.font.bold = True

    if "Resume Body" not in styles:
        body_style = styles.add_style("Resume Body", WD_STYLE_TYPE.PARAGRAPH)
        body_style.font.name = "Arial"
        body_style.font.size = Pt(10.5)


def build_document(lines: list[str]) -> Document:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    ensure_styles(document)

    title = lines[0].strip()
    contact = lines[1].strip() if len(lines) > 1 else ""

    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.name = "Arial"
    title_run.font.size = Pt(18)

    if contact:
        contact_paragraph = document.add_paragraph()
        contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_run = contact_paragraph.add_run(contact)
        contact_run.font.name = "Arial"
        contact_run.font.size = Pt(10.5)

    body_lines = lines[2:]

    for raw_line in body_lines:
        line = raw_line.rstrip()
        stripped = line.strip()

        if not stripped:
            continue

        if is_upper_heading(stripped):
            paragraph = document.add_paragraph(style="Resume Section")
            paragraph.paragraph_format.space_before = Pt(8)
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.add_run(stripped)
            continue

        if stripped.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.space_after = Pt(2)
            run = paragraph.add_run(stripped[2:])
            run.font.name = "Arial"
            run.font.size = Pt(10.5)
            continue

        if looks_like_role_or_company(stripped) or looks_like_date_or_subrole(stripped):
            paragraph = document.add_paragraph(style="Resume Role")
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.add_run(stripped)
            continue

        paragraph = document.add_paragraph(style="Resume Body")
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.add_run(stripped)

    return document


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate a DOCX resume from the plain text resume source.")
    parser.add_argument(
        "input",
        nargs="?",
        default="assets/rodrigo-marin-resume.txt",
        help="Path to the text resume source.",
    )
    parser.add_argument(
        "output",
        nargs="?",
        default="assets/rodrigo-marin-resume.docx",
        help="Path to the generated DOCX file.",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    input_path = Path(args.input)
    output_path = Path(args.output)

    lines = input_path.read_text(encoding="utf-8").splitlines()
    document = build_document(lines)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    print(f"Wrote {output_path}")


if __name__ == "__main__":
    main()